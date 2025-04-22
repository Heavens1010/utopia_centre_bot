
import os
from docx import Document as DocxDocument
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.schema import Document
from dotenv import load_dotenv

load_dotenv()

DOC_DIR = "docs"
VECTOR_DIR = "vector_store"

def load_docs(doc_dir):
    docs = []
    for fname in os.listdir(doc_dir):
        if fname.endswith(".docx"):
            path = os.path.join(doc_dir, fname)
            doc = DocxDocument(path)
            text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            docs.append(Document(page_content=text, metadata={"source": fname}))
    return docs

def build_vector_store():
    print("📥 Loading DOCX documents...")
    documents = load_docs(DOC_DIR)
    print(f"📚 {len(documents)} documents loaded.")

    embedding = OpenAIEmbeddings(openai_api_key=os.getenv("OPENAI_API_KEY"))
    db = Chroma.from_documents(documents, embedding, persist_directory=VECTOR_DIR)
    db.persist()
    print(f"✅ Vector store created in: {VECTOR_DIR}/")

if __name__ == "__main__":
    build_vector_store()
